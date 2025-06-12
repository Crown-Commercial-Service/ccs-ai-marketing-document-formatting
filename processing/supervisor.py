import os
import re
from difflib import SequenceMatcher
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from utils.blob_utils import upload_to_blob  # 👈 Import the blob uploader
from processing.word_replacement import replace_restricted_words  # ✅ Import word replacement function
from processing.tone import process_chunks  # Tone Adjustment Agent
from processing.preprocess import chunk_text, extract_text  # 🔥 Updated import
from processing.writing_guidelines import process_writing_guidelines # Writing Guidelines Agent
#from processing.acronym_expander import process_acronyms
from processing.acronym_expander import apply_acronym_expansion
from processing.formatting import process_formatting  # Formatting Agent # Preprocessing Agent
from processing.audit import audit_document_with_openai
import openai  # For LLM-assisted merging

import uuid

job_id = str(uuid.uuid4())[:8]  # Shorten for readability


# Folder where final processed files are stored
PROCESSED_FOLDER = os.path.abspath("backend/formatted_docs")
os.makedirs(PROCESSED_FOLDER, exist_ok=True)

CHUNK_SIZE = 200  # Words per chunk
OVERLAP = 20  # Words to overlap between chunks

# LLM Configuration
AZURE_OPENAI_ENDPOINT = os.getenv("AZURE_OPENAI_ENDPOINT")
AZURE_OPENAI_KEY = os.getenv("AZURE_OPENAI_KEY")
DEPLOYMENT_NAME = os.getenv("DEPLOYMENT_NAME")
API_VERSION = "2024-02-01"

client = openai.AzureOpenAI(
    azure_endpoint=AZURE_OPENAI_ENDPOINT,
    api_key=AZURE_OPENAI_KEY,
    api_version=API_VERSION
)


def calculate_similarity(original, processed):
    """Calculate similarity between original and processed text."""
    return SequenceMatcher(None, original, processed).ratio()


def review_chunks(original_chunks, processed_chunks, threshold=0.5):
    """Ensures processed chunks are not significantly altered."""
    approved_chunks = []

    for i, (orig, proc) in enumerate(zip(original_chunks, processed_chunks)):
        similarity = calculate_similarity(orig, proc)

        if similarity < threshold:
            print(f"⚠️ Warning: Chunk {i+1} has low similarity ({similarity:.2f}). Reviewing changes carefully.")
            approved_chunks.append(proc if len(proc) > len(orig) * 0.8 else orig)  # Use processed chunk unless it's too short
        else:
            approved_chunks.append(proc)  # Use processed version

    return approved_chunks

def merge_chunks_with_llm(chunks):
    """Uses LLM to merge and structure chunks intelligently while maintaining formatting consistency."""

    system_prompt = """
        You are an advanced document editor. Your job is to merge and format text chunks into a well-structured, coherent document.

        **Strict Rules to Follow:**
        - **DO NOT lose the content,DO NOT duplicate content. Ensure each point or paragraph appears only once across the entire document, even if similar wording exists in multiple chunks and remove unwanted space**.
        - **DO NOT split the bullet points further more, DO NOT alter the capitalization of bullet points.** Preserve their original case.
        - **Ensure whole document is in active voice and Ensure no "Chunk X" artifacts** appear in the final output.
        - **Put the most important information first, Use simple words—assume the reader has no prior knowledge,Replace complex phrases with single words where possible.**
        - **Do NOT lowercase words unnecessarily, Proper nouns, acronyms, and key terms should remain capitalized and ensure every words are in UK spelling**.
        - **Bullet points and numbered lists must be preserved as they are.**
        - **DO NOT use Oxford commas, semicolons in bullet points as well**
        - **Write all numbers in numerals except when they start a sentence, even for the numbers in the bullet points**
        - **Maintain logical flow and smooth transitions.**
        - **Use bold headers** for major sections.
        - **If a table format is detected (like agreements), keep it structured.**
        - **Do NOT add extra spaces before or after bullet points.**
        - **Do NOT insert "---" or "**" unless originally present.**
        - **Ensure headings retain their original formatting** (do not add extra `**`).
        - **DO NOT duplicate content when merging chunks**.
        - **Ensure statements from speakers (e.g., CEO quotes) remain properly formatted**.
        - **Except Heading, use "you" and "we" instead of "readers" and "CCS" respectively.**
        - Keep bullet points structured and **do not introduce new headings** above them.
        **Return only the fully merged text. Do NOT add explanations.**
    """

    # 🛠️ **Fix: Remove Duplicates Before Sending to LLM**
    unique_chunks = []
    seen_texts = set()
    for chunk in chunks:
        if chunk.strip() not in seen_texts:
            unique_chunks.append(chunk.strip())
            seen_texts.add(chunk.strip())

    response = client.chat.completions.create(
        model=DEPLOYMENT_NAME,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": "\n\n".join(unique_chunks)}
        ],
        temperature=0.0
    )

    merged_text = response.choices[0].message.content.strip()

    # 🔥 **Fix 1: Remove "Chunk X:" Artifacts**
    merged_text = re.sub(r"Chunk \d+:?", "", merged_text).strip()

    # 🔥 **Fix 2: Standardize Bullet Point Formatting**
    merged_text = re.sub(r"●", "-", merged_text)  # Convert "●" to "-"
    merged_text = re.sub(r"•", "-", merged_text)  # Convert "•" to "-"

    # 🔥 **Fix 3: Ensure Bullet Points Retain Proper Case and Spacing**
    def fix_bullet_points(match):
        bullet = match.group(1).strip()
        return f"\n- {bullet}"  # Ensure spacing is correct and case is preserved

    merged_text = re.sub(r"\n- ([A-Za-z].+)", fix_bullet_points, merged_text)

    # 🔥 **Fix 4: Ensure Numbered Lists Start with a Capital Letter**
    def fix_numbered_list(match):
        number, text = match.groups()
        return f"\n{number} {text.strip().capitalize()}"  # Capitalize first word in numbered lists

    merged_text = re.sub(r"\n(\d+\.)\s([a-z].+)", fix_numbered_list, merged_text)

    # 🔥 **Fix 5: Remove Extra Blank Lines After Bullet Points**
    merged_text = re.sub(r"\n(- .+?)\n\s*\n", r"\n\1\n", merged_text)

    # 🔥 **Fix 6: Ensure CEO Quotes & Statements Retain Proper Formatting**
    merged_text = re.sub(r'(?<=\n)“(.+?)”', r'\n"\1"', merged_text)  # Convert curly quotes to straight ones

    return merged_text


from odf.opendocument import OpenDocumentText
from odf.style import Style, TextProperties, ParagraphProperties
from odf.text import P, H, List, ListItem

def save_final_document(filename, merged_text,job_id):
    """Saves the final formatted document to a .docx file with proper styling."""
   
    processed_folder_path = os.path.join(PROCESSED_FOLDER, job_id)
    os.makedirs(processed_folder_path, exist_ok=True)
 
    final_save_path = os.path.join(processed_folder_path, filename)
    doc = Document()

    def add_heading(text, level=1):
        """Adds a properly formatted heading."""
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(14) if level == 1 else Pt(12)
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    def add_bullet(text):
        """Adds a bullet point while maintaining spacing."""
        para = doc.add_paragraph(style="List Bullet")
        run = para.add_run(text)
        run.font.size = Pt(11)

    def add_paragraph(text):
        """Adds a regular paragraph with spacing."""
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.size = Pt(11)
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Convert formatted text into structured paragraphs
    lines = merged_text.split("\n")
   
    for line in lines:
        line = line.strip()
       
        if line.startswith("**") and line.endswith("**"):  # Convert **Bold Headings**
            add_heading(line.replace("**", ""), level=1)        
        elif line.startswith("### "):  # Convert H2 headings
            add_heading(line.replace("### ", ""), level=2)
        elif line.startswith("## "):  # Convert H2 headings
            add_heading(line.replace("## ", ""), level=2)
        elif line.startswith("# "):  # Convert H1 headings
            add_heading(line.replace("# ", ""), level=1)
        elif line.startswith("- "):  # Convert bullet points
            add_bullet(line.replace("- ", ""))
        elif line.startswith("**") and line.endswith("**"):  # Convert **Bold Headings**
            add_heading(line.replace("**", ""), level=1)
        else:
            add_paragraph(line)

    doc.save(final_save_path)
    print(f"✅ Final document saved locally at: {final_save_path}")

    # Upload to Azure Blob Storage
    blob_path = f"formatted_docs/{job_id}/{filename}"
    upload_to_blob(final_save_path, blob_path)

    return blob_path  # 👈 This will now be used instead of local file path



def supervise_document(file_path, job_id):
    """
    Supervises document processing:
    - Extracts original text
    - Replaces restricted words
    - Chunks it for processing
    - Applies tone adjustments and verifies changes
    - Applies writing guidelines and verifies changes
    - Applies formatting rules and verifies all prior changes
    - Expands acronyms correctly across chunks
    - Merges chunks with LLM assistance
    - Saves formatted document
    """
    print(f"\n🛠 Step 1: Extracting and cleaning paragraphs from {file_path}")


    # Step 1: Extract text and chunk it
    paragraphs = extract_text(file_path)

    full_text = "\n".join(paragraphs)

    # Step 2: Apply restricted word replacement BEFORE chunking
    print("\n🔄 Applying restricted word replacements before chunking...")
    cleaned_text = replace_restricted_words(full_text)
    cleaned_paragraphs = cleaned_text.split("\n")
    print(f"🔧 Processing started for Job ID: {job_id}")


    # Step 3: Chunk text
    chunked_texts = chunk_text(cleaned_paragraphs, max_words=CHUNK_SIZE, overlap=OVERLAP)
    print("\n🔹 ORIGINAL CHUNKS BEFORE PROCESSING 🔹")
    for i, chunk in enumerate(chunked_texts):
        print(f"\nChunk {i+1}:\n{chunk}\n{'-'*50}")

    # Step 4: Apply tone adjustments
    tone_adjusted_chunks = process_chunks(chunked_texts)
    verified_tone_chunks = review_chunks(chunked_texts, tone_adjusted_chunks)

    # Step 5: Apply writing guidelines
    writing_guidelines_chunks = process_writing_guidelines(verified_tone_chunks)
    verified_writing_chunks = review_chunks(verified_tone_chunks, writing_guidelines_chunks)

    # Step 6: Apply formatting rules
    formatted_chunks = process_formatting(verified_writing_chunks)
    verified_formatted_chunks = review_chunks(verified_writing_chunks, formatted_chunks)

    # Step 7: Apply acronym logic after merging
    merged_text = merge_chunks_with_llm(verified_formatted_chunks)

    # Now apply acronym expansion on the entire merged text instead of chunks
    expanded_text = apply_acronym_expansion(merged_text)


    # Remove job_id prefix if it exists
    original_filename = os.path.basename(file_path)
    if original_filename.startswith(f"{job_id}_"):
        original_filename = original_filename[len(f"{job_id}_"):]
  # e.g., 'input.docx'
   
    print("original_filename",original_filename)

    # ✅ Use original filename for output too — clean and exact
    final_path = save_final_document(original_filename, expanded_text, job_id)

    # ✅ Also use original filename for audit report
    audit_path = audit_document_with_openai(full_text, expanded_text, job_id, original_filename)

    return final_path, audit_path
