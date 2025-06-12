import os
import re
import logging
from docx import Document
from processing.tone import process_chunks  # ✅ Correct import for Tone Adjustment Agent
from processing.framework_processor import FrameworkProcessor  # ✅ Import FrameworkProcessor
from odf.opendocument import load
from odf.text import P

PROCESSED_FOLDER = os.path.abspath("backend/formatted_docs")
os.makedirs(PROCESSED_FOLDER, exist_ok=True)

CHUNK_SIZE = 100  # Words per chunk
OVERLAP = 20  # Words to overlap between chunks

def extract_text_from_docx(file_path):
    """Extracts text from a .docx file while preserving paragraph structure."""
    doc = Document(file_path)
    paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    return paragraphs
def extract_text(file_path):
    """Extracts text from a .docx or .odt file."""
    if file_path.endswith(".docx"):
        doc = Document(file_path)
        paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
        return paragraphs

    elif file_path.endswith(".odt"):
        paragraphs = []
        doc = load(file_path)
        all_paragraphs = doc.getElementsByType(P)
        for para in all_paragraphs:
            text = ""
            for n in para.childNodes:
                if n.nodeType == 3:  # Text node
                    text += n.data
            if text.strip():
                paragraphs.append(text.strip())
        return paragraphs

    else:
        raise ValueError("Unsupported file format. Only .docx and .odt are supported.")


def chunk_text(paragraphs, max_words=CHUNK_SIZE, overlap=OVERLAP):
    """Splits text into chunks of approximately max_words with overlap while preserving structure."""
    words = []
    paragraph_map = []  # Tracks which words belong to which paragraph

    # Flatten paragraphs into words while tracking their original paragraph
    for para_index, para in enumerate(paragraphs):
        split_words = para.split()
        words.extend(split_words)
        paragraph_map.extend([para_index] * len(split_words))  # Map each word to its paragraph index

    chunks = []
    start = 0

    while start < len(words):
        end = min(start + max_words, len(words))

        # Extract chunked words only
        chunk_words = words[start:end]
        chunk_text = " ".join(chunk_words)  # Join only the words for this chunk

        chunks.append(chunk_text)

        start += max_words - overlap  # Move forward with overlap

    return chunks

def preprocess_document(file_path, framework_csv):
    """Extracts text, processes chunks through Tone Adjustment, and saves the final output."""
   
    # Step 1: Extract raw text from the document
    paragraphs = extract_text(file_path)
    full_text = "\n".join(paragraphs)

    # Step 2: Apply framework name replacement BEFORE chunking
    logging.info("🔄 Applying framework replacements before chunking...")
    framework_processor = FrameworkProcessor(framework_csv)
    processed_text = framework_processor.replace_frameworks(full_text)

    # Step 3: Normalize spaces in paragraphs
    cleaned_paragraphs = [' '.join(para.split()) for para in processed_text.split("\n") if para.strip()]

    # Step 4: Apply chunking
    chunked_texts = chunk_text(cleaned_paragraphs, max_words=CHUNK_SIZE, overlap=OVERLAP)

    # 🔥 Debugging: Print chunked text BEFORE tone adjustment
    print("\n🔹 CHUNKED TEXT BEFORE TONE ADJUSTMENT 🔹")
    for i, chunk in enumerate(chunked_texts, start=1):
        print(f"\nChunk {i}:\n{chunk}\n{'-'*50}")  # Separates chunks for readability

    # Step 5: Apply tone adjustment
    processed_chunks = process_chunks(chunked_texts)

    # Step 6: Save processed document
    base_filename = os.path.splitext(os.path.basename(file_path))[0]
    processed_filename = base_filename + ".docx"
    processed_path = os.path.join(PROCESSED_FOLDER, processed_filename)


    doc = Document()

    for i, chunk in enumerate(processed_chunks, start=1):
        doc.add_paragraph(f"Chunk {i}:")  # Add chunk header
        doc.add_paragraph(chunk)  # Add processed chunk content
        doc.add_paragraph("")  # Ensure spacing between chunks

    doc.save(processed_path)

    print(f"✅ Preprocessed and tone-adjusted document saved at: {processed_path}")
    return processed_path


