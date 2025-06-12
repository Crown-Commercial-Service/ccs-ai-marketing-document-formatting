import os

import openai

from docx import Document

from docx.shared import Pt, RGBColor

from docx.enum.text import WD_ALIGN_PARAGRAPH

from datetime import datetime
from utils.blob_utils import upload_to_blob  # ✅ Correct name

from dotenv import load_dotenv
# Load environment variables
load_dotenv()
# Load environment variables

AZURE_OPENAI_ENDPOINT = os.getenv("AZURE_OPENAI_ENDPOINT")

AZURE_OPENAI_KEY = os.getenv("AZURE_OPENAI_KEY")

DEPLOYMENT_NAME = os.getenv("DEPLOYMENT_NAME")

API_VERSION = "2024-02-01"
 
client = openai.AzureOpenAI(

    azure_endpoint=AZURE_OPENAI_ENDPOINT,

    api_key=AZURE_OPENAI_KEY,

    api_version=API_VERSION

)
 
AUDIT_FOLDER = os.path.abspath("backend/audit_reports")

os.makedirs(AUDIT_FOLDER, exist_ok=True)
 


RULES = """

1. Use plain English.

2. Write in short sentences.

3. Use active voice.

4. Avoid jargon and legalistic language.

5. Use correct headings and structure for easy navigation.

6. Avoid unnecessary formatting like underline or italics.

7. Provide meaningful link text (not "click here").

8. Always use sentence case for headings.

9. Use bold for headings (never underline or italicise).

10. Separate paragraphs with a blank line.

11. Bullet points must start with a lowercase letter unless it’s a proper noun.

12. No punctuation at the end of bullet points unless full sentences.

13. Do not use Oxford commas.

14. Avoid long bullet lists; split if necessary.

15. Maintain left alignment for all text (never justified).

16. Always write in UK English spelling.

17. Avoid American English (eg. "organisation" not "organization").

18. Avoid Latin or non-English words (eg. do not use "etc.", "eg.", "ie.").

19. Use simple, common words instead of complex alternatives.

20. Avoid exclamation marks (!).

21. Do not use semicolons.

22. Use “you” and “we” rather than “the user” or “Crown Commercial Service” in text.

23. Avoid passive voice.

24. Write numbers as numerals (1, 2, 3) not words (one, two, three) except at the start of sentences.

25. Use commas for numbers over 1,000 (e.g. 1,200).

26. Use % symbol for percentages (e.g. 10% not 10 percent).

27. Always round currency values sensibly unless exact figures are critical.

28. Do not mix number formats within a list.

29. For ranges, use "to" not hyphens (e.g. 10 to 20).

30. Avoid corporate jargon ("leverage", "synergy", etc.).

31. Use direct, clear instructions.

32. Replace vague words (e.g. "various", "several") with specific details.

33. Replace Jargon words, buzzwords with plain alternatives.

34. Avoid "world-class", "cutting-edge", "best practice" type phrases.

35. Use only UK English

36. Avoid contractions in formal content (use "do not" instead of "don't").

37. Avoid political or controversial words.

38. Avoid slang or overly casual language.

39. Always spell Crown Commercial Service fully on first use (then CCS thereafter).

40. Acronyms must be expanded fully at first mention.

41. Capitalise official names and services (e.g. "Digital Marketplace").

42. Maintain CCS brand tone: clear, confident, trustworthy.

43. Never use the CCS logo incorrectly.

44. Structure pages to be scannable: use headings, short paragraphs, bullet points.

45. Prioritise important information at the top.

46. Keep sentences and paragraphs short.

47. Use meaningful link labels.

48. Avoid duplicate content.

49. Do not refer to download sizes or formats unnecessarily (e.g. "PDF").

50. Expand acronyms in body text, even if expanded in headings.

51. Proper noun references must stay capitalised throughout.

52. Maintain consistency of tense throughout a document.

53. Only use quotation marks for direct speech or quotes.

54. Avoid overuse of bold text (only headings).

55. Do not start headings with numbers unless essential.

56. If listing items, use parallel sentence structures.

57. Maintain consistency in formatting (bullet styles, numbering, etc.)

58. Ensure all documents are clear for screen readers (logical reading order).

59. Use appropriate alt text for any included images (if applicable).

60. Tables must have headers and clear structure.

61. Prefer positive phrasing over negative where appropriate.

62. Focus on clarity, simplicity, and user needs at every point.

"""
def chunk_by_paragraph(text, max_chars=1800):

    paragraphs = text.split('\n\n')

    chunks = []

    current_chunk = ""
 
    for para in paragraphs:

        para = para.strip()

        if not para:

            continue
 
        # If adding this paragraph exceeds max, start a new chunk

        if len(current_chunk) + len(para) + 2 > max_chars:

            chunks.append(current_chunk.strip())

            current_chunk = para

        else:

            if current_chunk:

                current_chunk += "\n\n" + para

            else:

                current_chunk = para
 
    if current_chunk:

        chunks.append(current_chunk.strip())
 
    return chunks

 
def audit_document_with_openai(original_text, final_text, job_id, filename):

    system_prompt = f"""
You are an expert document quality auditor for Crown Commercial Services (CCS).

You will receive:

- Original document (before processing)
- Final document (after processing)

You must only describe missed rules and applied changes. Do not include a rule classification table.


You must:

- Create a clean, beautifully structured report.
- Include a clearly separated section titled "Rules Not Covered" and "Not in Scope" that explains exactly what was missed, and shows examples from the document where the rule was not followed.
- Add a Before vs After table listing every original vs transformed text snippet, rule number(s), and a clear explanation of the applied change. Each row should represent one change, not summaries.
- Include a paragraph-by-paragraph and line-by-line breakdown.
  - For every paragraph, list *all* rules it satisfies (if it satisfies 10+ rules, show all 10+). And explain how the rule has changed the input to output.
  - Each paragraph breakdown must cover all applicable rules in that paragraph. Do not skip any covered rule, even if the change seems minor.
  - Every rule listed as "✅ Covered" must appear in the paragraph-by-paragraph breakdown with clear evidence by comparing input and output.
  - provide a brief but clear justification of how each rule was fulfilled line by line .
  - Paragraph Breakdown should start with how many rules has been covered and how many changes has been made by rules in the output.
  - Paragraph Breakdown should be very clear and detailed, it should show specifically the changes made by all rules.
  - Paragraph Breakdown should have Input: , Output: . Where it should show how the rules has changed line by line inside paragraph has changed the input to output.
  - Paragraph Breakdown should not miss to explain each and every rules which has been applied.
  - Do not skip or group multiple rule coverages into one explanation. Each changes line by line must be treated separately.

STRICT FORMATTING INSTRUCTIONS:

- Do NOT use punctuation markdown unnecessarily (e.g. no `#`, `|`, `**`,'-|','--|')
- Use Heading1 for tables and headings (e.g. "Rules Not Covered", "Before vs After Table", etc.)
- Use simple numbered sections to separate parts of the report (e.g. 1. Summary Table, 2. Rules Not Covered, 3. Before vs After, 4. Paragraph Breakdown)
- Do NOT include general insights or summaries like "the document is overall good"
- Add space_after = Pt(6) to all normal text paragraphs to avoid crowding
- Do NOT invent new rules beyond the 62 listed
- Assume the reader has no knowledge of the document formatting task. Make it self-explanatory.
IMPORTANT: DO Not ignore any content by saying additionalAdditional paragraph breakdowns for all chunks would follow this format , List it for every paragraphs and content, line by line. Each rule listed as "✅ Covered" must appear in the Paragraph Breakdown with justification. If a rule does not appear there, consider it not covered and reclassify it accordingly. Do not skip rules or claim they are covered without concrete comparison evidence.

Here are the rules to audit against:

{RULES}
"""
    original_chunks = chunk_by_paragraph(original_text)
    final_chunks = chunk_by_paragraph(final_text)                      

    partial_audits = []
 
    for i, (orig_chunk, final_chunk) in enumerate(zip(original_chunks, final_chunks)):
        chunk_prompt = f"Original Document:\n{orig_chunk}\n\nFinal Document:\n{final_chunk}"
       
        response = client.chat.completions.create(
            model=DEPLOYMENT_NAME,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": chunk_prompt}
            ],
            temperature=0.0
        )

        content = response.choices[0].message.content.strip()
        if content:
            print(f"✅ Chunk {i+1}/{len(original_chunks)} processed")
            partial_audits.append(content)
        else:
            print(f"❌ Chunk {i+1}/{len(original_chunks)} returned empty")


 
    audit_output = "\n\n---\n\n".join(partial_audits)
 
 
    return save_audit_report(audit_output, job_id, filename)
 
def save_audit_report(audit_text, job_id, original_filename):
    doc = Document()

    # Cover Page
    heading = doc.add_heading('Document Formatting Audit Report', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Job ID: {job_id}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Filename: {original_filename}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Date: {datetime.now().strftime('%d %B %Y')}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Generated by: Crown Commercial Services Formatter System").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # Structured output processing
    sections = audit_text.split('---')
    for section in sections:
        lines = section.strip().split('\n')
        if not lines:
            continue

        title = lines[0].strip()
        if title:
            heading = doc.add_heading(title, level=1)
            heading.runs[0].font.color.rgb = RGBColor(0x00, 0x32, 0x96)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        for line in lines[1:]:
            if line.startswith("|") and "|" in line:
                cells = [cell.strip() for cell in line.split('|') if cell.strip()]
                if len(cells) >= 2:
                    if 'current_table' not in locals():
                        current_table = doc.add_table(rows=1, cols=len(cells))
                        hdr_cells = current_table.rows[0].cells
                        for i, cell_text in enumerate(cells):
                            hdr_cells[i].text = cell_text
                    else:
                        row_cells = current_table.add_row().cells
                        for i, cell_text in enumerate(cells):
                            row_cells[i].text = cell_text
            else:
                if 'current_table' in locals():
                    current_table.style = 'Table Grid'
                    del current_table
                paragraph = doc.add_paragraph(line)
                paragraph.style.font.size = Pt(11)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 🔥 Create subfolder for job_id to match Flask download path
    job_folder = os.path.join(AUDIT_FOLDER, job_id)
    os.makedirs(job_folder, exist_ok=True)

    audit_filename = f"audit_{original_filename}"
    audit_path = os.path.join(job_folder, audit_filename)
    doc.save(audit_path)
    print(f"✅ Audit Report saved locally at: {audit_path}")

    # Upload to Azure Blob Storage
    blob_path = f"audit_reports/{job_id}/audit_{original_filename}"
    upload_to_blob(audit_path, blob_path)

    return blob_path  # 👈 Return blob path instead of local path


